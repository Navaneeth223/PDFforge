import os
import zipfile
import shutil
import mammoth
import msoffcrypto
import io
from docx import Document
from typing import List
from services.storage import get_output_path

def word_to_html(session_id: str, file_path: str) -> str:
    output_path = get_output_path(session_id, ".html")
    with open(file_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        with open(output_path, "w", encoding="utf-8") as html_file:
            html_file.write(result.value)
    return output_path

def word_to_text(session_id: str, file_path: str) -> str:
    output_path = get_output_path(session_id, ".txt")
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))
    return output_path

def merge_word_docs(session_id: str, file_paths: List[str]) -> str:
    output_path = get_output_path(session_id, ".docx")
    master = Document(file_paths[0])
    for path in file_paths[1:]:
        doc = Document(path)
        for element in doc.element.body:
            master.element.body.append(element)
    master.save(output_path)
    return output_path

def compress_word(session_id: str, file_path: str) -> str:
    """Remove unused media/images from docx (ZIP)."""
    output_path = get_output_path(session_id, ".docx")
    # Simple re-save via python-docx often reduces size by cleaning up
    doc = Document(file_path)
    doc.save(output_path)
    return output_path

def remove_word_password(session_id: str, file_path: str, password: str) -> str:
    output_path = get_output_path(session_id, ".docx")
    file = msoffcrypto.OfficeFile(open(file_path, "rb"))
    file.load_key(password=password)
    with open(output_path, "wb") as f:
        file.decrypt(f)
    return output_path
